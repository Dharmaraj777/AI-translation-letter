import io
import textwrap
from typing import List, Dict, Optional

from docx import Document  # python-docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from PIL import Image, ImageDraw, ImageFont

from spanish_translator_logger import logger
from spanish_translator_oai_client import OaiClient


class DocxProcessor:
    """
    DOCX translator.

    - Text in the document body (paragraphs + tables) is translated at RUN level,
      preserving font, size, color, etc.
    - Images:
        * For each embedded image:
            - Send image bytes to GPT-4.1 vision
            - GPT reads & translates any text in the image
            - We render the translated text into a new image (same size, white background)
            - Replace the image in the DOCX with this translated version.
    """

    def __init__(self, oai_client: OaiClient):
        self.oai_client = oai_client


    # TEXT: collect RUN-level segments
    # ---------------------------------------------------------
    def _collect_segments(self, doc: Document) -> List[Dict[str, str]]:
        segments: List[Dict[str, str]] = []

        # Top-level paragraphs
        for p_idx, para in enumerate(doc.paragraphs):
            for r_idx, run in enumerate(para.runs):
                text = run.text
                if text and text.strip():
                    seg_id = f"p-{p_idx}-r-{r_idx}"
                    segments.append({"id": seg_id, "text": text})

        # Tables (cells)
        for t_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        for r_idx, run in enumerate(para.runs):
                            text = run.text
                            if text and text.strip():
                                seg_id = (
                                    f"tbl-{t_idx}-row-{row_idx}-col-{col_idx}-"
                                    f"p-{p_idx}-r-{r_idx}"
                                )
                                segments.append({"id": seg_id, "text": text})

        return segments

    def _apply_text_translations(
        self,
        doc: Document,
        id_to_translation: Dict[str, str],
    ) -> None:
        # Top-level paragraphs
        for p_idx, para in enumerate(doc.paragraphs):
            for r_idx, run in enumerate(para.runs):
                text = run.text
                if not text or not text.strip():
                    continue
                seg_id = f"p-{p_idx}-r-{r_idx}"
                if seg_id in id_to_translation:
                    run.text = id_to_translation[seg_id]

        # Tables
        for t_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        for r_idx, run in enumerate(para.runs):
                            text = run.text
                            if not text or not text.strip():
                                continue
                            seg_id = (
                                f"tbl-{t_idx}-row-{row_idx}-col-{col_idx}-"
                                f"p-{p_idx}-r-{r_idx}"
                            )
                            if seg_id in id_to_translation:
                                run.text = id_to_translation[seg_id]

    # ---------------------------------------------------------
    # IMAGES: markdown-table parser + GPT-4.1 vision + redraw
    # ---------------------------------------------------------
    def _parse_markdown_table(self, translated_text: str):
        """
        Very lightweight parser for Markdown-style tables that GPT returns, e.g.:

            Exemple de récompense

            | Points         | Montant      |
            |----------------|--------------|
            | 200 points     | 10 $         |
            | 500 points     | 25 $         |

        Returns:
            title_lines: list[str]  (lines before the first table row)
            rows: list[list[str]]   (each row is a list of cell strings)
        """
        lines = [l.rstrip() for l in translated_text.splitlines()]

        title_lines: list[str] = []
        table_lines: list[str] = []
        in_table = False

        for line in lines:
            stripped = line.strip()
            if stripped.startswith("|") and "|" in stripped[1:]:
                in_table = True
                if not stripped.endswith("|"):
                    stripped = stripped + "|"
                table_lines.append(stripped)
            else:
                if not in_table:
                    if stripped:
                        title_lines.append(stripped)

        if not table_lines:
            return title_lines, []

        parsed_rows = []
        for ln in table_lines:
            inner = ln.strip().strip("|")
            parts = [p.strip() for p in inner.split("|")]

            is_sep = all(not p or set(p) <= {"-", ":"} for p in parts)
            if is_sep:
                continue

            parsed_rows.append(parts)

        if not parsed_rows:
            return title_lines, []

        n_cols = len(parsed_rows[0])
        cleaned_rows = [row for row in parsed_rows if len(row) == n_cols]

        return title_lines, cleaned_rows

    def _render_translated_image(
        self,
        original_bytes: bytes,
        translated_text: str,
        content_type: str,
    ) -> bytes:
        """
        Create a new image same size as original, white background, and:
        - If GPT returned a Markdown-like table, draw a table (title + grid)
        - Otherwise, draw wrapped text
        """
        try:
            orig_img = Image.open(io.BytesIO(original_bytes)).convert("RGB")
        except Exception as e:
            logger.error(f"Failed to open original image for redraw: {e}")
            return original_bytes

        width, height = orig_img.size
        new_img = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(new_img)

        try:
            font = ImageFont.load_default()
        except Exception:
            font = None

        # Safe line height
        if font is not None:
            try:
                ascent, descent = font.getmetrics()
                line_height = ascent + descent + 4
            except Exception:
                line_height = 16
        else:
            line_height = 16

        margin = 20

        # Try table mode first
        title_lines, table_rows = self._parse_markdown_table(translated_text)

        if table_rows:
            logger.info("[image] Rendering translated image as table layout.")

            # 1) Draw title
            y = margin
            if title_lines:
                for line in title_lines:
                    if y + line_height > height - margin:
                        break
                    draw.text((margin, y), line, fill="black", font=font)
                    y += line_height + 4
                y += 8

            # 2) Column widths
            n_cols = len(table_rows[0])
            col_lengths = [0] * n_cols
            for row in table_rows:
                for j, cell in enumerate(row):
                    col_lengths[j] = max(col_lengths[j], len(cell))

            total_len = sum(col_lengths) or 1
            available_width = width - 2 * margin
            col_widths = []
            for L in col_lengths:
                w_j = max(int(available_width * (L / total_len)), 80)
                col_widths.append(w_j)

            scale = available_width / float(sum(col_widths))
            col_widths = [int(w * scale) for w in col_widths]

            table_top = y
            row_height = line_height + 10

            for row_idx, row in enumerate(table_rows):
                x = margin
                row_top = table_top + row_idx * row_height
                row_bottom = row_top + row_height

                if row_bottom > height - margin:
                    break

                for col_idx, cell_text in enumerate(row):
                    col_width = col_widths[col_idx]
                    cell_left = x
                    cell_right = x + col_width

                    draw.rectangle(
                        [cell_left, row_top, cell_right, row_bottom],
                        outline="black",
                        width=1,
                    )

                    text_x = cell_left + 5
                    text_y = row_top + (row_height - line_height) // 2

                    draw.text((text_x, text_y), cell_text, fill="black", font=font)

                    x += col_width

        else:
            # Fallback: plain wrapped text
            logger.info("[image] No markdown table detected; rendering as wrapped text.")
            wrapped_lines: List[str] = []
            for paragraph in translated_text.splitlines():
                if not paragraph.strip():
                    wrapped_lines.append("")
                    continue
                wrapped_lines.extend(textwrap.wrap(paragraph, width=60))

            y = margin
            for line in wrapped_lines:
                if y + line_height > height - margin:
                    break
                draw.text((margin, y), line, fill="black", font=font)
                y += line_height

        out = io.BytesIO()
        fmt = "PNG"
        if "jpeg" in content_type.lower() or "jpg" in content_type.lower():
            fmt = "JPEG"
        new_img.save(out, format=fmt)
        out.seek(0)
        return out.read()

    def _translate_images_in_doc(
        self,
        doc: Document,
        target_language: Optional[str] = None,
        target_dialect: Optional[str] = None,
    ) -> None:
        image_count = 0
        translated_count = 0

        for rel_id, rel in doc.part.rels.items():
            if rel.reltype != RT.IMAGE:
                continue

            image_count += 1
            image_part = rel.target_part
            original_bytes = image_part.blob
            content_type = getattr(image_part, "content_type", "image/png")

            logger.info(f"[image] Found image rel_id={rel_id}, content_type={content_type}")

            try:
                translated_text = self.oai_client.translate_image_to_language(
                    original_bytes,
                    content_type=content_type,
                    target_language=target_language,
                    target_dialect=target_dialect,
                )
            except Exception as e:
                logger.error(f"GPT-4.1 vision failed for image (rel_id={rel_id}): {e}")
                continue

            if not translated_text.strip():
                logger.info(f"[image] No translated text returned for rel_id={rel_id}; leaving image unchanged.")
                continue

            logger.info(
                f"[image] GPT translation for rel_id={rel_id} (first 80 chars): "
                f"{translated_text[:80]!r}"
            )

            try:
                new_bytes = self._render_translated_image(
                    original_bytes,
                    translated_text,
                    content_type,
                )
                # Use private _blob because .blob is read-only in your python-docx version
                image_part._blob = new_bytes
                translated_count += 1
                logger.info(f"[image] Replaced image (rel_id={rel_id}) with translated version.")
            except Exception as e:
                logger.error(f"Failed to render/replace translated image (rel_id={rel_id}): {e}")
                continue

        logger.info(f"[image] Completed image translation. Found={image_count}, translated={translated_count}")

    
    def translate_document(
        self,
        filename: str,
        content_bytes: bytes,
        target_language: Optional[str] = None,
        target_dialect: Optional[str] = None,
    ) -> bytes:
        logger.info(f"Translating DOCX document (text + images): {filename}")
        doc_stream = io.BytesIO(content_bytes)
        doc = Document(doc_stream)

        # 1) Text (runs)
        segments = self._collect_segments(doc)
        if segments:
            id_to_translation = self.oai_client.translate_segments(
                segments,
                target_language=target_language,
                target_dialect=target_dialect,
            )
            self._apply_text_translations(doc, id_to_translation)
        else:
            logger.info("No text segments found in DOCX body text.")

        # 2) Images via GPT-4.1 vision
        try:
            self._translate_images_in_doc(
                doc,
                target_language=target_language,
                target_dialect=target_dialect,
            )
        except Exception as e:
            logger.error(f"Image translation step failed (non-fatal): {e}")

        # 3) Save final DOCX
        out_stream = io.BytesIO()
        doc.save(out_stream)
        out_stream.seek(0)
        return out_stream.read()