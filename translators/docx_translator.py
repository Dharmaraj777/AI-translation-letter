# translators/docx_translator.py

import io
import textwrap
from typing import List, Dict, Optional

from docx import Document  # python-docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from PIL import Image, ImageDraw, ImageFont

from .base_translator import BaseTranslator
from ai_translation_logger import logger
from ai_translation_utils import UtilityFunctions


class DocxTranslator(BaseTranslator):
    """
    DOCX translator.

    - Text in the document body (paragraphs + tables) is translated at RUN level,
      preserving font, size, color, etc.
    - Images:
        * For each embedded image:
            - Send image bytes to GPT-4.1 vision
            - GPT reads & translates any text in the image
            - We render the translated text into a new image (same size, white background)
            - Replace the image in the DOCX with this translated version.
    """

    def can_handle(self, filename: str) -> bool:
        return UtilityFunctions.get_extension(filename) == ".docx"

    # ---------------------------------------------------------
    # TEXT: collect RUN-level segments
    # ---------------------------------------------------------
    def _collect_segments(self, doc: Document) -> List[Dict[str, str]]:
        segments: List[Dict[str, str]] = []

        # Top-level paragraphs
        for p_idx, para in enumerate(doc.paragraphs):
            for r_idx, run in enumerate(para.runs):
                text = run.text
                if text and text.strip():
                    seg_id = f"p-{p_idx}-r-{r_idx}"
                    segments.append({"id": seg_id, "text": text})

        # Tables (cells)
        for t_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        for r_idx, run in enumerate(para.runs):
                            text = run.text
                            if text and text.strip():
                                seg_id = (
                                    f"tbl-{t_idx}-row-{row_idx}-col-{col_idx}-"
                                    f"p-{p_idx}-r-{r_idx}"
                                )
                                segments.append({"id": seg_id, "text": text})

        return segments

    def _apply_text_translations(
        self,
        doc: Document,
        id_to_translation: Dict[str, str],
    ) -> None:
        # Top-level paragraphs
        for p_idx, para in enumerate(doc.paragraphs):
            for r_idx, run in enumerate(para.runs):
                text = run.text
                if not text or not text.strip():
                    continue
                seg_id = f"p-{p_idx}-r-{r_idx}"
                if seg_id in id_to_translation:
                    run.text = id_to_translation[seg_id]

        # Tables
        for t_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        for r_idx, run in enumerate(para.runs):
                            text = run.text
                            if not text or not text.strip():
                                continue
                            seg_id = (
                                f"tbl-{t_idx}-row-{row_idx}-col-{col_idx}-"
                                f"p-{p_idx}-r-{r_idx}"
                            )
                            if seg_id in id_to_translation:
                                run.text = id_to_translation[seg_id]

    # ---------------------------------------------------------
    # IMAGES: GPT-4.1 vision + redraw
    # ---------------------------------------------------------
    def _render_translated_image(
        self,
        original_bytes: bytes,
        translated_text: str,
        content_type: str,
    ) -> bytes:
        """
        Create a new image same size as original, white background,
        and draw the translated text with simple word wrapping.
        """
        try:
            orig_img = Image.open(io.BytesIO(original_bytes)).convert("RGB")
        except Exception as e:
            logger.error(f"Failed to open original image for redraw: {e}")
            return original_bytes

        width, height = orig_img.size
        new_img = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(new_img)

        try:
            font = ImageFont.load_default()
        except Exception:
            font = None

        margin = 20
        max_text_width_px = width - 2 * margin

        # Simple character-based wrapping (we don't know real font metrics here)
        wrapped_lines: List[str] = []
        for paragraph in translated_text.splitlines():
            if not paragraph.strip():
                wrapped_lines.append("")
                continue
            wrapped_lines.extend(textwrap.wrap(paragraph, width=60))

        if font is not None:
            line_height = font.getsize("A")[1] + 4
        else:
            line_height = 16

        y = margin
        for line in wrapped_lines:
            if y + line_height > height - margin:
                break
            draw.text((margin, y), line, fill="black", font=font)
            y += line_height

        out = io.BytesIO()
        fmt = "PNG"
        if "jpeg" in content_type.lower() or "jpg" in content_type.lower():
            fmt = "JPEG"
        new_img.save(out, format=fmt)
        out.seek(0)
        return out.read()

    def _translate_images_in_doc(
        self,
        doc: Document,
        target_language: Optional[str] = None,
        target_dialect: Optional[str] = None,
    ) -> None:
        """
        For each embedded image:
          - send to GPT-4.1 vision for translation
          - redraw image with translated text
          - replace image bytes in the DOCX part
        """
        image_count = 0
        translated_count = 0

        for rel_id, rel in doc.part.rels.items():
            if rel.reltype != RT.IMAGE:
                continue

            image_count += 1
            image_part = rel.target_part
            original_bytes = image_part.blob
            content_type = getattr(image_part, "content_type", "image/png")

            logger.info(f"[image] Found image rel_id={rel_id}, content_type={content_type}")

            try:
                translated_text = self.oai_client.translate_image_to_language(
                    original_bytes,
                    content_type=content_type,
                    target_language=target_language,
                    target_dialect=target_dialect,
                )
            except Exception as e:
                logger.error(f"GPT-4.1 vision failed for image (rel_id={rel_id}): {e}")
                continue

            if not translated_text.strip():
                logger.info(f"[image] No translated text returned for rel_id={rel_id}; leaving image unchanged.")
                continue

            logger.info(f"[image] GPT translation for rel_id={rel_id} (first 80 chars): {translated_text[:80]!r}")

            try:
                new_bytes = self._render_translated_image(
                    original_bytes,
                    translated_text,
                    content_type,
                )
                # IMPORTANT: use the public property .blob
                image_part.blob = new_bytes
                translated_count += 1
                logger.info(f"[image] Replaced image (rel_id={rel_id}) with translated version.")
            except Exception as e:
                logger.error(f"Failed to render/replace translated image (rel_id={rel_id}): {e}")
                continue

        logger.info(f"[image] Completed image translation. Found={image_count}, translated={translated_count}")

    # ---------------------------------------------------------
    # Public entrypoint
    # ---------------------------------------------------------
    def translate_document(
        self,
        filename: str,
        content_bytes: bytes,
        target_language: Optional[str] = None,
        target_dialect: Optional[str] = None,
    ) -> bytes:
        logger.info(f"Translating DOCX document (text + images): {filename}")
        doc_stream = io.BytesIO(content_bytes)
        doc = Document(doc_stream)

        # 1) Text (runs)
        segments = self._collect_segments(doc)
        if segments:
            id_to_translation = self.oai_client.translate_segments(
                segments,
                target_language=target_language,
                target_dialect=target_dialect,
            )
            self._apply_text_translations(doc, id_to_translation)
        else:
            logger.info("No text segments found in DOCX body text.")

        # 2) Images via GPT-4.1 vision
        try:
            self._translate_images_in_doc(
                doc,
                target_language=target_language,
                target_dialect=target_dialect,
            )
        except Exception as e:
            logger.error(f"Image translation step failed (non-fatal): {e}")

        # 3) Save final DOCX
        out_stream = io.BytesIO()
        doc.save(out_stream)
        out_stream.seek(0)
        return out_stream.read()
